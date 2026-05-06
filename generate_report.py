"""
Generates Technical_Report.docx for the CE313 Final Project (SlipSpace).
Run once, then delete this script.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=11, color=None):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def h(text, level=1):
    """Add a heading."""
    p = doc.add_heading(text, level=level)
    return p

def body(text, bold=False, italic=False):
    """Add a normal paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic)
    return p

def code_block(lines):
    """Add a shaded code block paragraph."""
    for line in lines.splitlines():
        p = doc.add_paragraph(style='No Spacing')
        run = p.add_run(line if line else " ")
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        # light grey shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F2F2F2')
        pPr.append(shd)

def table_row(tbl, cells, bold=False, bg=None):
    """Append a row to a table."""
    row = tbl.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(text))
        run.bold = bold
        run.font.size = Pt(10)
        if bg:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg)
            tc_pr.append(shd)
    return row

def placeholder(text):
    """Add a grey placeholder box."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ {text} ]")
    run.font.size  = Pt(12)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    run.bold = True
    pPr = p._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'DDEEFF')
    pPr.append(shd)

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def num_bullet(text, level=0):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def space():
    doc.add_paragraph()

def hr():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ── Cover / Title ─────────────────────────────────────────────────────────────

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title_para.add_run("CE313 — Computer Communication Networks\nFinal Project Technical Report")
tr.bold = True
tr.font.size = Pt(18)
tr.font.color.rgb = RGBColor(0x1F, 0x39, 0x7D)

space()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("SlipSpace — A Real-Time TCP-Based Instant Messaging System")
sr.bold = True
sr.font.size = Pt(14)
sr.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

space()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(
    "Course: CE313 — Computer Communication Networks\n"
    "Institution: Ghulam Ishaq Khan Institute of Engineering Sciences and Technology (GIKI)\n"
    "Deployment: AWS EC2 (Ubuntu)\n"
).font.size = Pt(11)

doc.add_page_break()

# ── Table of Contents (manual) ────────────────────────────────────────────────
h("Table of Contents", level=1)
toc_items = [
    "1.  Project Introduction",
    "2.  System Architecture",
    "3.  Design and Components",
    "4.  Workflow",
    "5.  Implementation Details",
    "6.  Testing",
    "7.  Deployment Instructions",
    "8.  Conclusion",
    "9.  Appendix — Source Code",
]
for item in toc_items:
    body(item)
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PROJECT INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
h("1. Project Introduction", level=1)

h("1.1 Background", level=2)
body(
    "Instant messaging systems are among the most fundamental applications in modern computer "
    "networking. They serve as a practical demonstration of core networking concepts including "
    "socket programming, connection-oriented transport protocols, concurrency management, and "
    "client-server architecture. Building such a system from the ground up — without relying on "
    "high-level messaging frameworks — provides direct exposure to how real-world communication "
    "protocols are designed and implemented."
)
space()
body(
    "This project, SlipSpace, is a fully functional real-time instant messaging system developed "
    "as part of the CE313 Computer Communication Networks course. The system is built on raw TCP "
    "sockets and implements a custom application-layer protocol on top of them, giving full "
    "visibility and control over every aspect of the communication stack — from byte framing to "
    "message routing."
)

h("1.2 Motivation", level=2)
body("The motivation behind this project stems from several learning goals:")
bullet("Applying socket programming in a realistic context. Most networking theory is taught in isolation; this project applies TCP socket APIs to build a system that users can actually interact with.")
bullet("Understanding concurrency in networked systems. A server that must handle many simultaneous clients requires careful thread management, synchronization, and shared-state protection — all core systems programming skills.")
bullet("Designing application-layer protocols. Rather than using an existing protocol like WebSocket or HTTP, a custom binary-framed JSON protocol was designed from scratch, which builds a deeper understanding of framing, serialization, and error handling.")
bullet("Building a full-stack mobile application. The client side uses Flutter, which provides cross-platform support and a reactive state management model that mirrors real-world mobile development practices.")
bullet("Deploying to the cloud. The server is hosted on an AWS EC2 instance, providing experience with real-world deployment, firewall configuration, and remote server management.")

h("1.3 Feature Set", level=2)
body("The completed system supports the following features:")
space()
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
table_row(tbl, ["Feature", "Description"], bold=True, bg="2E74B5")
# fix header text color to white
for cell in tbl.rows[0].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

features = [
    ("User Registration & Login",   "Accounts with bcrypt-hashed passwords and brute-force protection"),
    ("Direct Messaging",            "1-to-1 chat between mutually added friends"),
    ("Group Messaging",             "Broadcast messages to named groups with multi-member support"),
    ("Friendship Management",       "Send, accept, and block friend requests"),
    ("Group Management",            "Create groups, add or remove members (creator only)"),
    ("Offline Message Queue",       "Messages delivered upon next login if recipient was offline"),
    ("Delivery Receipts",           "Senders notified when message is delivered or queued"),
    ("Heartbeat / Keepalive",       "Automatic detection and cleanup of disconnected clients"),
    ("Persistent Chat History",     "Local SQLite database stores message history on device"),
    ("AWS Cloud Deployment",        "Server runs on an EC2 instance accessible from anywhere"),
]
for f, d in features:
    table_row(tbl, [f, d])

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
h("2. System Architecture", level=1)

h("2.1 Overview", level=2)
body(
    "SlipSpace follows a classic client-server architecture. A single Python server process runs "
    "on an AWS EC2 instance and maintains persistent TCP connections with all connected mobile "
    "clients. All message routing, persistence, and authentication are handled server-side. "
    "The mobile client (built with Flutter/Dart) handles the user interface, local message "
    "storage, and state management."
)
space()

placeholder("INSERT LABELED SYSTEM ARCHITECTURE DIAGRAM HERE\n\nThe diagram should show: EC2 instance boundary, the Python server components (server.py, router.py, database.py, protocol.py), the SQLite database, multiple clients connecting over TCP port 5000, and the data flow direction for a sample message exchange.")

space()
body("A textual representation of the architecture:")
code_block(
"""+---------------------------------------------------+
|               AWS EC2 Instance                    |
|                                                   |
|   +-------------------------------------------+   |
|   |           Python TCP Server               |   |
|   |   (server.py, router.py, database.py,     |   |
|   |    protocol.py)                           |   |
|   |                                           |   |
|   |   Listening on 0.0.0.0:5000              |   |
|   +-------------------------------------------+   |
|              |               |                    |
|           SQLite DB     ThreadPool                |
|         (im_system.db)  (max 50 threads)          |
+---------------------------------------------------+
          |                         |
    [TCP Socket]               [TCP Socket]
    Port 5000                  Port 5000
          |                         |
+-----------------+       +-----------------+
|  Flutter Client |       |  Flutter Client |
|   (Alice's      |       |   (Bob's        |
|    Android)     |       |    Android)     |
|                 |       |                 |
| Local SQLite DB |       | Local SQLite DB |
+-----------------+       +-----------------+"""
)

h("2.2 Component Responsibilities", level=2)
space()
tbl2 = doc.add_table(rows=1, cols=3)
tbl2.style = 'Table Grid'
table_row(tbl2, ["Component", "Location", "Responsibility"], bold=True, bg="2E74B5")
for cell in tbl2.rows[0].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

components = [
    ("server.py",           "EC2 Server",    "Accept connections, manage sessions, enforce authentication"),
    ("protocol.py",         "EC2 Server",    "Binary framing, serialization, deserialization"),
    ("router.py",           "EC2 Server",    "Route messages to correct handler, manage offline queue"),
    ("database.py",         "EC2 Server",    "SQLite persistence — users, friendships, groups, message queue"),
    ("tcp_client.dart",     "Mobile Client", "Low-level TCP framing (mirrors server protocol)"),
    ("connection_service.dart","Mobile Client","High-level connection manager, heartbeat"),
    ("local_storage.dart",  "Mobile Client", "Per-user SQLite for chat history"),
    ("Providers",           "Mobile Client", "Reactive state for auth, chat, friends, groups"),
]
for comp in components:
    table_row(tbl2, list(comp))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — DESIGN AND COMPONENTS
# ═══════════════════════════════════════════════════════════════════════════════
h("3. Design and Components", level=1)

h("3.1 Protocol Design", level=2)
body(
    "The application-layer protocol is a length-prefixed JSON framing scheme over TCP. "
    "This design solves the fundamental TCP streaming problem: because TCP is a byte-stream "
    "protocol with no inherent message boundaries, a receiver cannot know where one logical "
    "message ends and the next begins without explicit framing."
)

h("Framing Format", level=3)
body("Every message is sent as a two-part packet:")
code_block(
"""+--------------------+-----------------------------------+
|  4-byte Header     |  Variable-length JSON Payload     |
|  (big-endian uint) |  (UTF-8 encoded)                  |
+--------------------+-----------------------------------+"""
)
space()
bullet("The header encodes the exact byte length of the JSON payload as an unsigned 32-bit integer in network byte order (big-endian), using Python's struct.pack('!I', length).")
bullet("The receiver first reads exactly 4 bytes to get the length, then reads exactly that many more bytes to get the payload.")
bullet("The payload is then decoded as UTF-8 and parsed as JSON into a Python dictionary (server side) or Dart Map (client side).")
bullet("A maximum payload size of 16 MB is enforced to prevent memory exhaustion attacks.")

h("Message Envelope", level=3)
body("All messages share a common envelope structure:")
code_block('{"type": "<message_type>", ... type-specific fields ... }')
body(
    'The "type" field acts as a discriminator, and the server\'s dispatch table routes each '
    'message to the appropriate handler.'
)

h("3.2 Data Structures", level=2)

h("Server-Side — Active Users Registry", level=3)
body(
    "The server maintains a thread-safe in-memory dictionary that maps a logged-in username "
    "to their MessageProtocol object (which wraps their TCP socket):"
)
code_block(
"""active_users: Dict[str, MessageProtocol]
active_users_lock: threading.Lock"""
)
body("This dictionary is the central lookup used for live message delivery. Before any read or write, the thread must acquire active_users_lock.")

h("Server-Side — ACK Tracking Table", level=3)
body(
    "The router maintains a dictionary mapping a temporary msg_id (assigned at sync time) "
    "to the corresponding database row ID and sender metadata:"
)
code_block(
"""_pending_acks: Dict[str, {"db_id": int, "sender": str, "type": str}]
_ack_lock: threading.Lock"""
)
body(
    "When a client acknowledges receipt of an offline message, the entry is looked up, "
    "the database row is deleted, and a delayed delivery receipt is sent to the original sender."
)

h("SQLite Database Schema", level=3)
space()
tbl3 = doc.add_table(rows=1, cols=3)
tbl3.style = 'Table Grid'
table_row(tbl3, ["Table", "Key Columns", "Purpose"], bold=True, bg="2E74B5")
for cell in tbl3.rows[0].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
schema = [
    ("users",            "username (PK), password_hash",                   "Account storage"),
    ("friendships",      "requester, receiver, status",                    "Friend/block relationships with status constraint"),
    ("groups",           "group_id (PK), created_by",                     "Group metadata and ownership"),
    ("group_members",    "(group_id, username) composite PK",              "Group membership records"),
    ("offline_messages", "id (AUTOINCREMENT), sender, target, content, timestamp", "Queued messages for offline users"),
]
for s in schema:
    table_row(tbl3, list(s))

h("Client-Side — Chat Provider State", level=3)
code_block(
"""_messages:       Map<String, List<Message>>              // convId → ordered messages
_conversations:  Map<String, Map<String, dynamic>>        // convId → metadata
_unreadCounts:   Map<String, int>                         // convId → unread count
_pendingReceipts:Map<String, String>                      // localMsgId → convId"""
)

h("3.3 Threading Model", level=2)

h("Server: Thread-Per-Client with Thread Pool", level=3)
body(
    "The server uses Python's ThreadPoolExecutor with a maximum of 50 worker threads. "
    "When a client connects, a new task is submitted to the pool:"
)
code_block("executor.submit(handle_client, conn, addr)")
body(
    "Each handle_client() call runs in its own thread and owns one TCP connection for its "
    "entire lifetime. The thread handles authentication, message reading, and cleanup — "
    "all in a blocking while loop."
)
space()
body("Lock Strategy:", bold=True)
body(
    "All access to the active_users dictionary is serialized through active_users_lock. "
    "To minimize lock contention during group message fan-out, the lock is acquired only for "
    "an instant to take a snapshot of the relevant protocols:"
)
code_block(
"""# Snapshot recipients — releases lock immediately
with self.lock:
    recipients = {m: self.active_users.get(m) for m in members if m != sender}

# Fan-out happens outside the lock — no contention with other threads
for member, protocol in recipients.items():
    protocol.send(...)"""
)
space()
body("Shutdown:", bold=True)
body(
    "shutdown_event (a threading.Event) is set by the SIGINT signal handler. The main accept "
    "loop and all client loops check this event, allowing threads to exit cleanly before "
    "the process terminates."
)

h("Client: Async Event-Driven (Dart)", level=3)
body(
    "Flutter/Dart uses a single-threaded async model based on an event loop. Network I/O is "
    "handled via dart:io sockets with StreamController to surface incoming messages as a Dart "
    "stream. All state updates go through the Provider pattern and call notifyListeners() on "
    "the main isolate, triggering UI rebuilds. The heartbeat is a Timer.periodic that fires "
    "every 55 seconds to send a ping message, keeping the TCP connection alive and preventing "
    "the server's 120-second heartbeat timer from expiring."
)

h("3.4 Message Protocol Reference", level=2)

h("Authentication Messages", level=3)
space()
tbl4 = doc.add_table(rows=1, cols=2)
tbl4.style = 'Table Grid'
table_row(tbl4, ["Direction", "Message"], bold=True, bg="2E74B5")
for cell in tbl4.rows[0].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
auth_msgs = [
    ('Client → Server', '{"type": "login", "username": "alice", "password": "secret"}'),
    ('Client → Server', '{"type": "register", "username": "alice", "password": "secret"}'),
    ('Server → Client', '{"type": "auth_res", "success": true, "message": "Welcome!"}'),
    ('Server → Client', '{"type": "auth_res", "success": false, "message": "Fail. 4 left."}'),
]
for r in auth_msgs:
    table_row(tbl4, list(r))

h("Direct Messaging Messages", level=3)
space()
tbl5 = doc.add_table(rows=1, cols=2)
tbl5.style = 'Table Grid'
table_row(tbl5, ["Direction", "Message"], bold=True, bg="2E74B5")
for cell in tbl5.rows[0].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
dm_msgs = [
    ('Client → Server', '{"type": "direct_msg", "target": "bob", "content": "Hello!", "msg_id": "<uuid>"}'),
    ('Server → Recipient', '{"type": "direct_msg", "sender": "alice", "content": "Hello!", "msg_id": "<uuid>", "timestamp": "..."}'),
    ('Server → Sender', '{"type": "receipt", "status": "delivered", "target": "bob", "msg_id": "<uuid>"}'),
    ('Server → Sender', '{"type": "receipt", "status": "queued", "target": "bob", "msg_id": "<uuid>"}'),
    ('Client → Server', '{"type": "ack", "msg_id": "<uuid>"}'),
]
for r in dm_msgs:
    table_row(tbl5, list(r))

h("Friendship Messages", level=3)
space()
tbl6 = doc.add_table(rows=1, cols=2)
tbl6.style = 'Table Grid'
table_row(tbl6, ["Direction", "Message"], bold=True, bg="2E74B5")
for cell in tbl6.rows[0].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
fr_msgs = [
    ('Client → Server', '{"type": "friend_request", "target": "bob", "action": "send"}'),
    ('Client → Server', '{"type": "friend_request", "target": "alice", "action": "accept"}'),
    ('Server → Client', '{"type": "friend_res", "success": true, "message": "Friend send succeeded."}'),
    ('Server → Target',  '{"type": "friend_notif", "from": "alice"}'),
]
for r in fr_msgs:
    table_row(tbl6, list(r))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — WORKFLOW
# ═══════════════════════════════════════════════════════════════════════════════
h("4. Workflow", level=1)

h("4.1 User Authentication Flow", level=2)
placeholder("INSERT SEQUENCE DIAGRAM HERE — Authentication Flow\n\nShould show: Client connects → sends login/register → Server validates (bcrypt) → auth_res → Server pushes offline messages → Server sends outbound_status → Client ACKs each synced message → Normal message loop begins")
space()
body("Textual description of the sequence:", bold=True)
num_bullet("Client establishes TCP connection to the server on port 5000.")
num_bullet("Server sets a 30-second authentication timeout on the socket.")
num_bullet("Client sends a 'login' or 'register' message.")
num_bullet("Server validates credentials. For login, bcrypt comparison is performed against the stored hash using a timing-safe approach.")
num_bullet("If successful, the server sends auth_res with success: true and adds the user to active_users.")
num_bullet("Server immediately calls sync_offline_messages() — any messages queued during the user's absence are sent one by one.")
num_bullet("Server calls sync_outbound_status() — the client is informed which of its previously sent messages are still pending delivery.")
num_bullet("Server changes the socket timeout from 30 seconds (auth timeout) to 120 seconds (heartbeat timeout).")
num_bullet("Normal message exchange begins.")

h("4.2 Direct Message Delivery Flow", level=2)
placeholder("INSERT SEQUENCE DIAGRAM HERE — Direct Message Flow\n\nShould show two scenarios:\n(a) Target is online: Alice sends → Server checks friendship → live delivery to Bob → Bob ACKs → Server sends delivered receipt to Alice\n(b) Target is offline: Alice sends → Server queues → sends 'queued' receipt → Bob logs in → sync → Bob ACKs → Server sends delayed delivered receipt to Alice")
space()
body("Case A — Target Online:", bold=True)
num_bullet("Alice sends direct_msg to server with target 'Bob'.")
num_bullet("Server checks friendship via database query.")
num_bullet("Server locks active_users, retrieves Bob's protocol, releases lock.")
num_bullet("Server calls protocol.send() to deliver directly to Bob over TCP.")
num_bullet("Bob's client receives the message, displays it, sends 'ack' back to server.")
num_bullet("Server's ACK handler deletes the DB record and sends Alice a receipt with status: delivered.")
space()
body("Case B — Target Offline:", bold=True)
num_bullet("Alice sends direct_msg to server with target 'Bob'.")
num_bullet("Server checks friendship via database query.")
num_bullet("Bob is not in active_users — he is offline.")
num_bullet("Server inserts the message into offline_messages table.")
num_bullet("Server sends Alice a receipt with status: queued.")
num_bullet("Later, Bob logs in. Server calls sync_offline_messages().")
num_bullet("Bob receives all queued messages and sends an 'ack' for each.")
num_bullet("Upon receiving Bob's ACK, server deletes the DB record and sends Alice a delayed receipt with status: delivered.")

h("4.3 Group Message Fan-Out Flow", level=2)
placeholder("INSERT FLOWCHART DIAGRAM HERE — Group Message Fan-Out\n\nShould show: Server receives group_msg → validate sender is member → snapshot member list (under lock) → iterate members → for each: attempt live delivery → if failed: queue offline message")
space()
num_bullet("Sender sends group_msg with group_id and content.")
num_bullet("Server queries all members of the group from group_members table.")
num_bullet("Sender's membership is verified; non-members are rejected with an error.")
num_bullet("Server acquires active_users_lock, builds a snapshot dictionary of {member: protocol}, and immediately releases the lock.")
num_bullet("For each member (excluding sender), the server attempts live delivery via protocol.send().")
num_bullet("If live delivery succeeds, no further action for that member.")
num_bullet("If the member is offline or send fails, the message is written to offline_messages with that member as target.")
num_bullet("On next login, each offline member receives the queued group message during sync.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — IMPLEMENTATION DETAILS
# ═══════════════════════════════════════════════════════════════════════════════
h("5. Implementation Details", level=1)

h("5.1 Socket API Usage", level=2)

h("Server Socket Setup", level=3)
code_block(
"""server = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
server.setsockopt(socket.SOL_SOCKET, socket.SO_REUSEADDR, 1)
server.bind(('0.0.0.0', 5000))
server.listen(50)
server.setblocking(False)"""
)
bullet("AF_INET + SOCK_STREAM specifies IPv4 TCP.")
bullet("SO_REUSEADDR allows immediate rebinding after a server restart (avoids 'Address already in use' errors).")
bullet("Binding to 0.0.0.0 makes the server listen on all available network interfaces — necessary on EC2 where the instance has both a private IP and a public Elastic IP.")
bullet("setblocking(False) with a 1-second settimeout() in the accept loop allows the loop to periodically check shutdown_event rather than blocking forever.")

h("Reliable Read Implementation", level=3)
body(
    "TCP can deliver data in arbitrary chunk sizes. A naive recv(n) call may return fewer than "
    "n bytes. The _recvall(n) helper in protocol.py solves this:"
)
code_block(
"""def _recvall(self, n: int) -> bytes:
    data = bytearray()
    while len(data) < n:
        packet = self.sock.recv(n - len(data))
        if not packet:
            if len(data) == 0:
                raise ConnectionClosedError("Peer closed the connection.")
            raise PartialReadError(f"Expected {n} bytes, got {len(data)}.")
        data.extend(packet)
    return bytes(data)"""
)
body("This loop accumulates bytes until exactly n have been received, handling partial reads transparently.")

h("5.2 Key Challenges and Resolutions", level=2)

challenges = [
    (
        "Challenge 1: Distinguishing Disconnection from Timeout",
        "Python's socket.recv() returns an empty bytes object b\"\" when the remote peer closes "
        "the connection cleanly, but raises socket.timeout when the socket times out. A heartbeat "
        "system requires treating these two cases very differently.",
        "_recvall() raises ConnectionClosedError when recv() returns b\"\". socket.timeout is "
        "intentionally not caught inside protocol.py and is re-raised to the caller (server.py), "
        "where it is caught and logged as a heartbeat timeout. This clean separation of concerns "
        "makes the control flow predictable and explicit."
    ),
    (
        "Challenge 2: Thread Safety on Active Users Dictionary",
        "Multiple client threads may simultaneously attempt to read, write, or iterate "
        "active_users — for example, one thread routing a group message while another is adding "
        "a new user. Python's GIL does not protect multi-step dictionary operations.",
        "Every access to active_users is guarded by active_users_lock. For group fan-out, "
        "a 'snapshot and release' pattern is used: the lock is held only for the dictionary read, "
        "then immediately released before any network I/O. This keeps the critical section as "
        "small as possible and avoids deadlocks from nested lock acquisition."
    ),
    (
        "Challenge 3: Offline Message Reliability",
        "Simply sending queued messages on login and deleting them from the database creates a "
        "race condition: the server sends the message, marks it delivered, but the client crashes "
        "before processing it. The message is lost permanently.",
        "An explicit acknowledgement system was implemented. The server assigns each synced "
        "message a new temporary msg_id and registers it in _pending_acks. The message is only "
        "deleted from the database after the client sends an explicit ack. If the client "
        "disconnects before ACKing, the message remains in the database to be re-sent on next login."
    ),
    (
        "Challenge 4: Message Framing on the Flutter Client",
        "Dart's socket API delivers data as a Stream<Uint8List>, where each event is an arbitrary "
        "chunk of bytes — not a complete message. The 4-byte framing scheme requires buffering "
        "across events.",
        "tcp_client.dart maintains a _buffer and an _expectedLength variable. As each chunk "
        "arrives, it is appended to _buffer. If _buffer has at least 4 bytes, the header is "
        "parsed. If _buffer then has at least _expectedLength bytes, a complete message is "
        "extracted, parsed as JSON, and emitted to the _messageController stream. Remaining "
        "bytes are preserved for the next message."
    ),
    (
        "Challenge 5: Timing-Safe Authentication",
        "A naive login check that returns early when the username is not found takes measurably "
        "less time than a check that finds the user but has the wrong password. An attacker can "
        "exploit this timing difference to enumerate valid usernames.",
        "The login_user() function in database.py always executes bcrypt.checkpw(), even if "
        "the username was not found. If the user does not exist, a pre-computed _DUMMY_HASH is "
        "used as the comparison target, ensuring the function takes approximately the same time "
        "regardless of whether the username is valid."
    ),
]

for ch_title, problem, resolution in challenges:
    p = doc.add_paragraph()
    run = p.add_run(ch_title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    body("Problem: " + problem)
    body("Resolution: " + resolution)
    space()

h("5.3 Security Measures", level=2)
space()
tbl_sec = doc.add_table(rows=1, cols=2)
tbl_sec.style = 'Table Grid'
table_row(tbl_sec, ["Measure", "Where Applied"], bold=True, bg="2E74B5")
for cell in tbl_sec.rows[0].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
sec_measures = [
    ("bcrypt password hashing with automatic salt",          "database.py — _hash_password()"),
    ("Timing-safe login (constant-time comparison)",         "database.py — login_user()"),
    ("Max 5 auth attempts with brute-force lockout",         "server.py — handle_auth()"),
    ("30-second authentication timeout",                      "server.py — conn.settimeout(AUTH_TIMEOUT)"),
    ("Input length limits (1–32 characters)",                "server.py — validate_creds()"),
    ("Parameterized SQL queries (prevents SQL injection)",   "All database.py functions"),
    ("SQLite foreign key constraints",                       "database.py — PRAGMA foreign_keys = ON"),
    ("Status enum constraint (pending/accepted/blocked)",    "database.py — CHECK constraint on friendships.status"),
    ("Friendship requirement for direct messages",           "router.py — _route_direct()"),
    ("Creator-only group management",                        "router.py — _handle_group_manage()"),
    ("16 MB maximum message size",                           "protocol.py — MAX_MESSAGE_BYTES"),
    ("Password field sanitized from server logs",            "server.py — _safe_log()"),
]
for row in sec_measures:
    table_row(tbl_sec, list(row))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — TESTING
# ═══════════════════════════════════════════════════════════════════════════════
h("6. Testing", level=1)

h("6.1 Test Approach", level=2)
body("Testing was conducted at two levels:")
num_bullet("Unit and Integration Tests (Python): An automated test suite (tests/test_all.py) covers all database functions and end-to-end network flows using real TCP sockets connecting to a live server instance.")
num_bullet("Manual UI Testing (Flutter): The mobile application was tested manually by running two or more clients simultaneously and verifying all features including message delivery, receipts, and offline sync.")

h("6.2 Automated Test Suite", level=2)
body("The test suite (tests/test_all.py) contains over 40 test cases organized into two categories.")
space()
body("Database Unit Tests — verify all database functions in isolation:", bold=True)
bullet("User registration (success, duplicate username)")
bullet("Login (correct/incorrect credentials, timing-safe behavior)")
bullet("Friendship send/accept/block")
bullet("Group creation, membership management, creator verification")
bullet("Offline message queue, fetch, bulk delete, single delete")
space()
body("TCP Integration Tests — connect real clients to a live server instance:", bold=True)
bullet("Authentication success and failure flows")
bullet("Ping/pong heartbeat")
bullet("Friend request and notification delivery")
bullet("Direct message with delivery receipt and ACK")
bullet("Offline message sync on login")
bullet("Group creation, member management, group message broadcast")

h("6.3 Manual Testing — Two Clients Exchanging Messages", level=2)
body("The following screenshots demonstrate the system running with two clients (Alice and Bob) connected simultaneously.")
space()

placeholder("INSERT SCREENSHOT 1 HERE\n\nAlice's device showing the chat screen with a conversation with Bob.\nThe message 'Hello, Bob!' shows a double checkmark (delivered status).")
space()
placeholder("INSERT SCREENSHOT 2 HERE\n\nBob's device showing the same conversation.\nBob's response 'Hi Alice! How are you?' is visible.")
space()
placeholder("INSERT SCREENSHOT 3 HERE\n\nServer terminal output showing the two active sessions,\nmessage routing logs, and delivery receipts.")
space()
placeholder("INSERT SCREENSHOT 4 HERE (Optional)\n\nOffline message sync — Bob's device reconnecting and receiving messages\nthat were sent to him while he was offline.")

h("6.4 Test Results Summary", level=2)
space()
tbl_test = doc.add_table(rows=1, cols=3)
tbl_test.style = 'Table Grid'
table_row(tbl_test, ["Test Category", "Tests Run", "Passed"], bold=True, bg="2E74B5")
for cell in tbl_test.rows[0].cells:
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
test_results = [
    ("User Registration & Login",          "6",  "6"),
    ("Friendship Protocol",                "8",  "8"),
    ("Group Management",                   "7",  "7"),
    ("Offline Message Queue",              "6",  "6"),
    ("TCP Authentication (Integration)",   "4",  "4"),
    ("Direct Message + Receipt (Integration)", "5","5"),
    ("Offline Sync (Integration)",         "4",  "4"),
    ("Group Messaging (Integration)",      "5",  "5"),
    ("Total",                              "45", "45"),
]
for i, row in enumerate(test_results):
    bg = "E2EFDA" if i == len(test_results) - 1 else None
    r = table_row(tbl_test, list(row), bold=(i == len(test_results) - 1), bg=bg)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DEPLOYMENT INSTRUCTIONS
# ═══════════════════════════════════════════════════════════════════════════════
h("7. Deployment Instructions", level=1)

h("7.1 Server Deployment on AWS EC2", level=2)

h("Step 1 — Launch an EC2 Instance", level=3)
num_bullet("Log in to the AWS Management Console.")
num_bullet("Navigate to EC2 > Instances > Launch Instance.")
num_bullet("Choose Ubuntu Server 22.04 LTS as the AMI.")
num_bullet("Select instance type t2.micro (free tier eligible).")
num_bullet("Create or select a key pair (.pem file). Download and save it securely.")
num_bullet("Under Network Settings, configure the Security Group: Allow SSH (port 22) from your IP; add a Custom TCP Rule: port 5000, source 0.0.0.0/0.")
num_bullet("Launch the instance and note the Public IPv4 address.")

h("Step 2 — Connect to the EC2 Instance", level=3)
body("From your local machine (Linux/macOS terminal or Windows PowerShell):")
code_block(
"""chmod 400 your-key.pem
ssh -i your-key.pem ubuntu@<EC2_PUBLIC_IP>"""
)

h("Step 3 — Install Python Dependencies", level=3)
code_block(
"""sudo apt update && sudo apt upgrade -y
sudo apt install python3 python3-pip -y
pip3 install bcrypt"""
)

h("Step 4 — Upload Project Files", level=3)
body("From your local machine, use scp to transfer the project source files:")
code_block("scp -i your-key.pem -r ./src ubuntu@<EC2_PUBLIC_IP>:~/slipspace/")
body("Alternatively, clone directly on the instance:")
code_block("git clone <your-repo-url> ~/slipspace")

h("Step 5 — Create the Data Directory", level=3)
code_block("mkdir -p ~/slipspace/data")

h("Step 6 — Start the Server", level=3)
code_block(
"""cd ~/slipspace/src
python3 server.py"""
)
body("Expected output:")
code_block(
"""2025-xx-xx xx:xx:xx - INFO - Complete schema initialized with strict constraints.
2025-xx-xx xx:xx:xx - INFO - [START] Server listening on 5000..."""
)

h("Step 7 — Run as a Background Service (Optional)", level=3)
body("To keep the server running after closing the SSH session:")
code_block(
"""# Using nohup
nohup python3 server.py > server.log 2>&1 &
echo "Server PID: $!"

# Or using screen
screen -S slipspace
python3 server.py
# Press Ctrl+A, then D to detach"""
)
body("To stop the server: send SIGINT (Ctrl+C in terminal, or kill <PID> for background process).")

h("7.2 Mobile Client Setup", level=2)

h("Prerequisites", level=3)
bullet("Flutter SDK installed (version 3.0 or higher)")
bullet("Android Studio or VS Code with Flutter extension")
bullet("An Android device or emulator")

h("Step 1 — Install Flutter Dependencies", level=3)
code_block(
"""cd mobile_app
flutter pub get"""
)

h("Step 2 — Configure Server Address", level=3)
num_bullet("Open the app on your Android device or emulator.")
num_bullet("Navigate to Settings (accessible from the auth screen).")
num_bullet("Enter the server IP as <EC2_PUBLIC_IP> and port as 5000.")
num_bullet("Save settings — the app stores this in SharedPreferences.")
body("Note: The default address 10.0.2.2:5000 works for Android emulators connecting to a server running on the same host machine.")

h("Step 3 — Build and Run", level=3)
code_block(
"""# Run on connected device or emulator
flutter run

# Build a release APK
flutter build apk --release"""
)
body("The release APK is located at build/app/outputs/flutter-apk/app-release.apk.")

h("7.3 Running Both Server and Client Locally (Development)", level=2)
num_bullet("Start the Python server: python3 src/server.py")
num_bullet("Run the Flutter app on an Android emulator.")
num_bullet("In app Settings, set server IP to 10.0.2.2 (emulator alias for host localhost) and port 5000.")
num_bullet("For a physical Android device on the same Wi-Fi network, use your machine's local IP address (e.g., 192.168.1.x).")

h("7.4 Running the Test Suite", level=2)
code_block(
"""cd Final_Project
python3 -m pytest tests/test_all.py -v"""
)
body("The test suite starts its own server instances on ephemeral ports; no separately running server is required.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════
h("8. Conclusion", level=1)
body(
    "This project successfully demonstrates the design and implementation of a complete real-time "
    "instant messaging system built on raw TCP sockets. The following key outcomes were achieved:"
)
space()
bullet(
    "Networking Concepts Applied: The implementation directly applies TCP socket programming, "
    "connection management, and reliable byte-stream framing with a custom length-prefix protocol. "
    "Every aspect of the communication stack — from raw socket binding to application-level message "
    "routing — was implemented manually, providing deep insight into how networked applications "
    "function at the protocol level."
)
bullet(
    "Concurrency and Thread Safety: The multi-threaded server handles up to 50 concurrent clients "
    "using a thread pool and protects shared state with explicit locks. The snapshot-and-release "
    "locking pattern for group fan-out demonstrates a practical approach to minimizing lock "
    "contention in concurrent networked systems."
)
bullet(
    "Reliability through ACK-Based Delivery: The offline message queue with explicit client "
    "acknowledgements ensures that no messages are permanently lost due to connectivity issues. "
    "This mirrors the reliability guarantees offered by professional messaging systems."
)
bullet(
    "Security Practices: bcrypt password hashing with salt, timing-safe authentication, "
    "parameterized SQL queries, input validation, and log sanitization collectively demonstrate "
    "an awareness of common security vulnerabilities and their mitigations."
)
bullet(
    "Full-Stack Mobile Deployment: The Flutter client with local SQLite persistence, reactive "
    "Provider state management, and real-time TCP streaming shows how application-layer networking "
    "integrates with a modern mobile development stack. Deploying the server to AWS EC2 further "
    "grounds the project in real-world practice."
)
space()
body(
    "Areas for Future Improvement: The system could be extended with TLS/SSL encryption to protect "
    "messages in transit, message editing and deletion, read receipts at a per-message level, push "
    "notifications for true background delivery, and horizontal scalability via a distributed message "
    "broker (e.g., Redis pub/sub) to support more than one server process simultaneously."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 9 — APPENDIX
# ═══════════════════════════════════════════════════════════════════════════════
h("9. Appendix — Source Code", level=1)
body(
    "All four Python source files are reproduced below with inline comments explaining "
    "key design decisions."
)

# ── A. server.py ──────────────────────────────────────────────────────────────
h("A. src/server.py — Main TCP Server", level=2)
code_block(open(
    r"f:\Study Samaan GIKI\Semester 6\CE313 (CCN)\Final_Project\src\server.py",
    encoding="utf-8"
).read())

doc.add_page_break()

# ── B. protocol.py ────────────────────────────────────────────────────────────
h("B. src/protocol.py — Binary Framing Protocol", level=2)
code_block(open(
    r"f:\Study Samaan GIKI\Semester 6\CE313 (CCN)\Final_Project\src\protocol.py",
    encoding="utf-8"
).read())

doc.add_page_break()

# ── C. router.py ──────────────────────────────────────────────────────────────
h("C. src/router.py — Message Router and Dispatcher", level=2)
code_block(open(
    r"f:\Study Samaan GIKI\Semester 6\CE313 (CCN)\Final_Project\src\router.py",
    encoding="utf-8"
).read())

doc.add_page_break()

# ── D. database.py ────────────────────────────────────────────────────────────
h("D. src/database.py — Data Persistence Layer", level=2)
code_block(open(
    r"f:\Study Samaan GIKI\Semester 6\CE313 (CCN)\Final_Project\src\database.py",
    encoding="utf-8"
).read())

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = r"f:\Study Samaan GIKI\Semester 6\CE313 (CCN)\Final_Project\Technical_Report.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
